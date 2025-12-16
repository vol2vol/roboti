import os
import sys
import zipfile
from datetime import datetime
from typing import List, Dict, Tuple, Optional

import pandas as pd

# Word
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont



def _to_int(s: str) -> int:
    return int(s.strip())


def _to_float(s: str) -> float:
    return float(s.replace(',', '.').strip())


def input_rows() -> pd.DataFrame:
    print("Введите данные о товарах/позициях. Минимум 5 строк.")
    print("Для каждой позиции укажите: Наименование | Количество | Стоимость(за единицу)")
    print("Пример: Флешка 32ГБ | 3 | 12,49")
    print("Чтобы закончить (после 5+ строк), оставьте 'Наименование' пустым и нажмите Enter.\n")

    rows: List[Dict] = []
    i = 1
    while True:
        name = input(f"[{i}] Наименование: ").strip()
        if name == "":
            if len(rows) >= 5:
                break
            else:
                print("Нужно минимум 5 строк. Продолжайте ввод.")
                continue

        qty_s = input(f"[{i}] Количество (целое): ").strip()
        price_s = input(f"[{i}] Стоимость за единицу (число, можно с запятой): ").strip()

        try:
            qty = _to_int(qty_s)
            price = _to_float(price_s)
            if qty < 0 or price < 0:
                raise ValueError
        except Exception:
            print("Ошибка: количество — целое >= 0, стоимость — число >= 0. Повторите ввод позиции.\n")
            continue

        rows.append({
            "Наименование": name,
            "Количество": qty,
            "Стоимость": round(price, 2)
        })
        i += 1
        print("— позиция добавлена.\n")

    df = pd.DataFrame(rows, columns=["Наименование", "Количество", "Стоимость"])
    df["Сумма"] = (df["Количество"] * df["Стоимость"]).round(2)
    return df



def save_excel_with_chart(df: pd.DataFrame, filepath: str, title: str = "Гистограмма по стоимости") -> None:
    with pd.ExcelWriter(filepath, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Данные", index=False)

        workbook = writer.book
        worksheet = writer.sheets["Данные"]

        n_rows = len(df)
        cat_range = f"=Данные!$A$2:$A${n_rows + 1}"  # Наименование
        val_range = f"=Данные!$C$2:$C${n_rows + 1}"  # Стоимость

        chart = workbook.add_chart({'type': 'column'})
        chart.add_series({
            'name': 'Стоимость',
            'categories': cat_range,
            'values': val_range,
            'data_labels': {'value': True},
        })
        chart.set_title({'name': title})
        chart.set_x_axis({'name': 'Наименование'})
        chart.set_y_axis({'name': 'Стоимость'})
        chart.set_legend({'position': 'bottom'})

        worksheet.insert_chart('E2', chart, {'x_scale': 1.1, 'y_scale': 1.1})


def save_word_report(df: pd.DataFrame, filepath: str, report_title: str) -> None:

    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run(report_title)
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"

    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)

    for i in range(rows):
        for j in range(cols):
            table.cell(i + 1, j).text = str(df.iloc[i, j])

    doc.add_paragraph()
    total = df["Сумма"].sum().round(2)
    p_total = doc.add_paragraph()
    run_total = p_total.add_run(f"Итоговая сумма: {total}")
    run_total.bold = True
    run_total.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    doc.save(filepath)



def _pick_existing(paths: List[str]) -> Optional[str]:
    for p in paths:
        if os.path.exists(p):
            return p
    return None


def _get_system_font_paths() -> Tuple[str, str]:

    candidates: List[Tuple[List[str], List[str]]] = []

    if os.name == "nt":
        win = r"C:\Windows\Fonts"
        candidates += [
            ([f"{win}\\arial.ttf"], [f"{win}\\arialbd.ttf"]),
            ([f"{win}\\calibri.ttf"], [f"{win}\\calibrib.ttf"]),
            ([f"{win}\\segoeui.ttf"], [f"{win}\\segoeuib.ttf"]),
            ([f"{win}\\verdana.ttf"], [f"{win}\\verdanab.ttf"]),
        ]

    for normals, bolds in candidates:
        n = _pick_existing(normals)
        b = _pick_existing(bolds)
        if n and b:
            return n, b

    raise RuntimeError(
    )


def _register_system_fonts(font_name: str = "ReportFont") -> Tuple[str, str]:

    normal_path, bold_path = _get_system_font_paths()
    pdfmetrics.registerFont(TTFont(font_name, normal_path))
    pdfmetrics.registerFont(TTFont(f"{font_name}-Bold", bold_path))
    pdfmetrics.registerFontFamily(
        font_name,
        normal=font_name,
        bold=f"{font_name}-Bold",
        italic=font_name,
        boldItalic=f"{font_name}-Bold",
    )
    return font_name, f"{font_name}-Bold"


def save_pdf(df: pd.DataFrame, filepath: str, report_title: str) -> None:

    font_regular, font_bold = _register_system_fonts()

    doc = SimpleDocTemplate(
        filepath, pagesize=A4,
        rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TitleCenter',
        parent=styles['Title'],
        fontName=font_bold,
        alignment=1,
        fontSize=18,
        leading=22,
        spaceAfter=12
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['BodyText'],
        fontName=font_regular,
        fontSize=11,
        leading=14
    )

    elements = [Paragraph(report_title, title_style), Spacer(1, 12)]

    data = [list(df.columns)] + df.values.tolist()
    data = [[(f"{x:.2f}" if isinstance(x, float) else str(x)) for x in row] for row in data]

    table = Table(data, hAlign='LEFT')
    table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, 0), font_bold),
        ('FONT', (0, 1), (-1, -1), font_regular),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#eeeeee")),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(table)

    total = float(df["Сумма"].sum().round(2))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f"<b>Итоговая сумма: {total:.2f}</b>", body_style))

    doc.build(elements)


def make_zip(artifact_paths: List[str], zip_path: str) -> None:
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in artifact_paths:
            zf.write(p, arcname=os.path.basename(p))



def main():
    # 1) Ввод данных
    df = input_rows()

    # 2) Имена файлов
    today = datetime.now().date()
    date_str = today.isoformat()
    base_name = f"Отчет_{date_str}"

    excel_path = f"{base_name}.xlsx"
    word_path = f"{base_name}.docx"
    pdf_path = f"{base_name}.pdf"
    zip_path = f"{base_name}.zip"

    report_title = f"Отчет от {date_str}"

    # 3) Excel
    try:
        save_excel_with_chart(df, excel_path, title="Гистограмма по колонке «Стоимость»")
        print(f"✔ Создан Excel: {excel_path}")
    except Exception as e:
        print(f"Ошибка при создании Excel: {e}", file=sys.stderr)
        return

    # 4) Word
    try:
        save_word_report(df, word_path, report_title)
        print(f"✔ Создан Word: {word_path}")
    except Exception as e:
        print(f"Ошибка при создании Word: {e}", file=sys.stderr)
        return

    # 5) PDF (системный TTF)
    try:
        save_pdf(df, pdf_path, report_title)
        print(f"✔ Создан PDF: {pdf_path}")
    except Exception as e:
        print(f"Ошибка при создании PDF: {e}", file=sys.stderr)
        return

    # 6) ZIP
    try:
        make_zip([excel_path, word_path, pdf_path], zip_path)
        print(f"✔ Создан ZIP: {zip_path}")
    except Exception as e:
        print(f"Ошибка при упаковке ZIP: {e}", file=sys.stderr)
        return

    print("\nГотово! Файлы созданы в текущей папке.")


if __name__ == "__main__":
    main()
